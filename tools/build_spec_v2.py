from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DIAG = Path('docs/diagramas')
OUT = Path('02_Especificacao_Requisitos_UML_DER_Safe_Student.docx')


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tcPr = cell._tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for name, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{name}'))
        if node is None:
            node = OxmlElement(f'w:{name}'); tcMar.append(node)
        node.set(qn('w:w'), str(value)); node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); node = OxmlElement('w:tblHeader'); node.set(qn('w:val'), 'true'); trPr.append(node)


def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr(); trPr.append(OxmlElement('w:cantSplit'))


def style_table(table, widths=None, font_size=9):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = 'Table Grid'
    for ri, row in enumerate(table.rows):
        prevent_row_split(row)
        if ri == 0: set_repeat_table_header(row)
        for ci, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(cell)
            if ri == 0: set_cell_shading(cell, 'D9E2F3')
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
                for r in p.runs:
                    r.font.name = 'Times New Roman'; r.font.size = Pt(font_size); r.bold = (ri == 0)
            if widths and ci < len(widths): cell.width = Cm(widths[ci])


def add_table(doc, headers, rows, widths=None, font_size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers): t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row): cells[i].text = str(value)
    style_table(t, widths, font_size); return t


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}'); p.add_run(text); return p


def add_body(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25); p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_after = Pt(6)
    p.add_run(text); return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3); p.add_run(text); return p


def add_caption(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text); r.font.size = Pt(10); r.italic = True


def add_image(doc, name, caption, width=9.3):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(DIAG / name), width=Inches(width)); add_caption(doc, caption)


def set_portrait(section):
    section.orientation = WD_ORIENT.PORTRAIT; section.page_width = Cm(21); section.page_height = Cm(29.7)
    section.top_margin = Cm(3); section.left_margin = Cm(3); section.right_margin = Cm(2); section.bottom_margin = Cm(2)


def set_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE; section.page_width = Cm(29.7); section.page_height = Cm(21)
    section.top_margin = Cm(2); section.left_margin = Cm(2); section.right_margin = Cm(2); section.bottom_margin = Cm(2)


def header_footer(section):
    hp = section.header.paragraphs[0]; hp.text = 'Safe Student - Especificação de Requisitos e Modelagem V2.0'; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in hp.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(9)
    fp = section.footer.paragraphs[0]; fp.text = 'Práticas Extensionistas Integradoras VI | Prof. Altemar Sales de Oliveira'; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in fp.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(9)


doc = Document(); set_portrait(doc.sections[0]); header_footer(doc.sections[0])
doc.styles['Normal'].font.name = 'Times New Roman'; doc.styles['Normal'].font.size = Pt(12); doc.styles['Normal'].paragraph_format.line_spacing = 1.5
for level in range(1, 4):
    st = doc.styles[f'Heading {level}']; st.font.name = 'Times New Roman'; st.font.bold = True; st.font.size = Pt(14 if level == 1 else 12)
    st.paragraph_format.space_before = Pt(12); st.paragraph_format.space_after = Pt(6)

for _ in range(2): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run('CURSO DE ENGENHARIA DE SOFTWARE'); r.bold = True; r.font.size = Pt(14)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run('PRÁTICAS EXTENSIONISTAS INTEGRADORAS VI').bold = True
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run('SAFE STUDENT'); r.bold = True; r.font.size = Pt(18)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run('ESPECIFICAÇÃO DE REQUISITOS, UML E MODELO RELACIONAL'); r.bold = True; r.font.size = Pt(16)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Revisão técnica V2.0 - requisitos verificáveis e diagramas com notação formal')
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Saquarema/RJ - 2026')
doc.add_page_break()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run('IDENTIFICAÇÃO DO DOCUMENTO').bold = True
add_table(doc, ['Campo', 'Informação'], [
    ('Integrantes', 'Flavio Gabrig Ferreira - 202323361; Richardson Conceição Ferreira - 202323181'),
    ('Gerente do projeto', 'Richardson Conceição Ferreira - 202323181'),
    ('Professor', 'Altemar Sales de Oliveira'), ('Versão', 'V2.0'), ('Data', '18 de agosto de 2026'), ('Local', 'Saquarema/RJ')], [4, 12], 10)
add_body(doc, 'Esta versão substitui a modelagem anterior de requisitos/UML/DER e deve ser usada como referência técnica para a apresentação e para a rastreabilidade com o código.')
doc.add_page_break()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run('SUMÁRIO').bold = True
for item in [
    '1 VISÃO GERAL ............................................................ 4',
    '2 ATORES E FRONTEIRAS DE RESPONSABILIDADE ................ 4',
    '3 CRITÉRIOS DE QUALIDADE DOS REQUISITOS ..................... 5',
    '4 REQUISITOS FUNCIONAIS ................................................ 6',
    '5 REQUISITOS NÃO FUNCIONAIS ......................................... 7',
    '6 REGRAS DE NEGÓCIO ....................................................... 9',
    '7 CASOS DE USO TEXTUAIS ................................................ 9',
    '8 DIAGRAMAS DE CASOS DE USO ....................................... 12',
    '9 DIAGRAMAS DE CLASSES ................................................ 16',
    '10 MODELO RELACIONAL (DER) ........................................... 19',
    '11 DIAGRAMA DE SEQUÊNCIA .............................................. 21',
    '12 ARQUITETURA IMPLEMENTADA ....................................... 22',
    '13 RASTREABILIDADE .......................................................... 23',
    '14 RESULTADO DA REVISÃO TÉCNICA .................................. 24',
    'REFERÊNCIAS ....................................................................... 25']:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2); p.add_run(item)
doc.add_page_break()

add_heading(doc, '1 VISÃO GERAL', 1)
add_body(doc, 'Este documento descreve o comportamento efetivamente demonstrável do Safe Student. A revisão V2.0 abandona requisitos genéricos ou aspiracionais e registra somente comportamentos existentes no MVP, metas não funcionais verificáveis e limitações assumidas. O objetivo é permitir que um professor, analista ou desenvolvedor consiga conferir cada requisito contra o código, os testes e a interface sem depender de interpretação subjetiva.')
add_heading(doc, '1.1 Escopo do MVP', 2)
for text in ['Autenticação e autorização por perfil.', 'Cadastro de aluno e vínculo com responsável pela Gestão Escolar.', 'Registro de entrada e saída por token/QR representado na demonstração.', 'Notificação interna aos responsáveis vinculados.', 'Histórico, relatório, exportação CSV e trilha de auditoria.', 'Comunicação entre perfis autorizados.', 'Coleta estruturada de feedback da validação acadêmica.']:
    add_bullet(doc, text)
add_heading(doc, '1.2 Fora do escopo', 2)
for text in ['Biometria, reconhecimento facial, NFC/RFID e catracas físicas.', 'Geolocalização contínua.', 'SMS, WhatsApp ou push real.', 'Banco de dados transacional de produção.', 'Certificação formal de conformidade com LGPD, WCAG ou OWASP.', 'Uso de dados reais sem autorização institucional e governança.']:
    add_bullet(doc, text)
add_body(doc, 'A aplicação é um MVP acadêmico com dados sintéticos. A arquitetura futura pode adotar outros componentes, mas esses componentes não são desenhados nem descritos como se já existissem.')

add_heading(doc, '2 ATORES E FRONTEIRAS DE RESPONSABILIDADE', 1)
add_table(doc, ['Ator', 'Responsabilidade no MVP', 'O que não pode fazer'], [
    ('Responsável', 'Acompanhar alunos vinculados, consultar histórico/notificações, gerar relatório do próprio escopo, comunicar-se com perfis permitidos e registrar feedback.', 'Registrar presença, cadastrar aluno, criar vínculo, consultar auditoria ou exportar validação.'),
    ('Portaria', 'Registrar entrada/saída de alunos ativos e comunicar-se com perfis permitidos.', 'Cadastrar aluno/vínculo, consultar auditoria ou acessar mensagens de terceiros.'),
    ('Gestão Escolar', 'Cadastrar aluno, criar vínculos, registrar presença, gerar relatórios, consultar auditoria, comunicar-se e administrar a demonstração.', 'Ler conversas privadas das quais não participa; tratar dados reais sem governança.'),
    ('Administrador técnico', 'Existe como perfil suportado no backend para tarefas técnicas, mas não é ator essencial da apresentação acadêmica.', 'Não deve ser apresentado como funcionalidade essencial do MVP.')], [3.5, 7.5, 6], 9)

add_heading(doc, '3 CRITÉRIOS DE QUALIDADE DOS REQUISITOS', 1)
add_body(doc, 'Cada requisito foi reescrito para obedecer a quatro critérios: uma obrigação por item, linguagem direta, possibilidade de verificação e correspondência com o software atual. Expressões vagas como “sistema completo”, “seguro”, “rápido” ou “compatível com LGPD” não são usadas como critério de aceite sem indicar como a afirmação seria comprovada.')
add_table(doc, ['Regra de escrita', 'Aplicação nesta versão'], [
    ('Obrigação explícita', 'Os RF usam “O sistema deve...”.'),
    ('Critério observável', 'Cada RF possui critério de aceite verificável.'),
    ('Sem implementação inventada', 'Banco relacional, hardware e serviços externos aparecem apenas como evolução futura.'),
    ('RNF mensurável', 'Quando uma meta ainda não possui evidência, o documento marca a evidência como pendente em vez de declarar atendimento.'),
    ('Rastreabilidade', 'RFs são ligados a rotas/funções e testes existentes.')], [5, 12], 10)

add_heading(doc, '4 REQUISITOS FUNCIONAIS', 1)
rf = [
('RF-01','O sistema deve autenticar usuário ativo por e-mail e senha.','Credenciais válidas criam sessão; inválidas retornam erro e não criam sessão.','Alta'),
('RF-02','O sistema deve autorizar cada operação no servidor conforme o perfil do usuário.','Uma chamada direta à API com perfil sem permissão é bloqueada com status de autorização apropriado.','Alta'),
('RF-03','A Gestão Escolar deve poder cadastrar aluno ativo informando nome, matrícula e turma válidos.','Cadastro inválido é rejeitado; matrícula e token não podem colidir com outro aluno.','Alta'),
('RF-04','A Gestão Escolar deve poder vincular um responsável ativo a um aluno ativo.','Após o vínculo, o responsável passa a receber o aluno em seu escopo de consulta.','Alta'),
('RF-05','Portaria e Gestão Escolar devem poder registrar entrada de aluno ativo por token válido.','A entrada é persistida com horário, operador e origem; segunda entrada pendente no mesmo dia é bloqueada.','Alta'),
('RF-06','Portaria e Gestão Escolar devem poder registrar saída somente após uma entrada válida no mesmo dia escolar.','Saída sem entrada anterior é bloqueada; saída válida é persistida.','Alta'),
('RF-07','Após um registro de presença, o sistema deve gerar notificação interna para cada responsável vinculado ao aluno.','A quantidade de responsáveis notificados é retornada e as notificações são gravadas para os respectivos usuários.','Alta'),
('RF-08','O sistema deve permitir consulta do histórico de presença apenas dos alunos pertencentes ao escopo do usuário autenticado.','Responsável vê somente vinculados; Portaria/Gestão veem alunos ativos necessários às suas funções.','Alta'),
('RF-09','O sistema deve gerar resumo de presença por aluno permitido ao usuário.','Resumo apresenta entradas, saídas, taxa de demonstração e último registro sem incluir aluno fora do escopo.','Média'),
('RF-10','O sistema deve exportar o histórico permitido em arquivo CSV autenticado.','Sem sessão a exportação é negada; com sessão o CSV contém apenas registros permitidos.','Média'),
('RF-11','O sistema deve permitir troca de mensagens somente entre perfis autorizados e exibir a conversa apenas aos participantes.','Manipular o identificador do destinatário não contorna a regra; Gestão não lê conversa privada de terceiros.','Média'),
('RF-12','O sistema deve registrar eventos críticos em trilha de auditoria e permitir consulta pela Gestão.','Login, presença, cadastro, vínculo, mensagens, feedback e restauração da demo produzem eventos com usuário e horário quando aplicável.','Alta'),
('RF-13','O usuário deve poder consultar suas próprias notificações e marcá-las como lidas.','Somente a notificação pertencente ao usuário autenticado pode ser alterada.','Média'),
('RF-14','O sistema deve coletar feedback estruturado da validação e permitir à Gestão exportar apenas evidências coletadas durante a apresentação.','Registros DEMO_SEED não entram nos indicadores nem no CSV de evidências coletadas.','Média'),
('RF-15','A Gestão Escolar deve poder restaurar a base de demonstração para o estado inicial.','A restauração recompõe o seed, registra evento de auditoria e invalida as sessões em memória.','Baixa')]
add_table(doc, ['ID','Descrição','Critério de aceite','Prioridade'], rf, [1.5,7,7,2], 8.5)

add_heading(doc, '5 REQUISITOS NÃO FUNCIONAIS', 1)
rnf = [
('RNF-01','Segurança de credenciais','Senhas de demonstração devem permanecer armazenadas em hash scrypt; comparação deve usar timingSafeEqual.','Testes de security.js.'),
('RNF-02','Sessão','A sessão deve possuir expiração configurável; padrão do MVP: 2 horas.','Inspeção do código e teste de autenticação.'),
('RNF-03','Privacidade por minimização','Diretórios e consultas devem retornar apenas os dados necessários; dados sintéticos devem ser usados na demonstração.','Testes de regressão de privacidade + inspeção da seed.'),
('RNF-04','Acessibilidade','A interface deve manter navegação por teclado, foco visível, rótulos e estrutura semântica compatíveis com avaliação de WCAG 2.2 quando aplicável.','Inspeção manual; não declarar certificação WCAG.'),
('RNF-05','Desempenho','No conjunto de dados da demonstração, respostas de uso corrente devem ter meta de até 2 s em ambiente local.','Meta a medir e anexar; não há evidência formal de carga neste documento.'),
('RNF-06','Compatibilidade de runtime','O backend deve executar em Node.js 18, 20 e 22.','GitHub Actions executa sintaxe e testes nessas versões.'),
('RNF-07','Portabilidade','O MVP não deve exigir banco, mensageria ou serviço externo para iniciar localmente.','npm start com módulos nativos do Node.'),
('RNF-08','Integridade de presença','O sistema deve impedir sequência inconsistente de entrada/saída e considerar o fuso escolar configurado ao definir o dia.','Testes de domínio/API.'),
('RNF-09','Integridade cadastral','Matrícula e token de aluno devem ser únicos no conjunto de dados.','Validação de matrícula + geração de token com verificação de colisão.'),
('RNF-10','Auditabilidade','Eventos críticos devem registrar identificador, usuário, ação e data/hora.','Endpoint de auditoria + testes.'),
('RNF-11','Testabilidade','Regras de domínio, segurança e API devem possuir testes automatizados executados no CI.','npm test e workflow ci.yml.'),
('RNF-12','Manutenibilidade','Regras reutilizáveis devem permanecer separadas entre domínio, segurança, servidor, interface e dados.','Estrutura Safe_Student_MVP/lib, server.js, public e tests.')]
add_table(doc, ['ID','Categoria','Requisito verificável','Evidência'], rnf, [1.5,3.5,7.5,5.5], 8.2)

add_heading(doc, '6 REGRAS DE NEGÓCIO', 1)
rules = [
('RN-01','Somente aluno ATIVO pode receber registro de presença.'),('RN-02','Entrada duplicada sem saída anterior no mesmo dia escolar deve ser recusada.'),('RN-03','Saída exige entrada anterior válida no mesmo dia escolar.'),('RN-04','Responsável consulta somente aluno explicitamente vinculado.'),('RN-05','Somente Gestão/Admin cadastra aluno e cria vínculo.'),('RN-06','Somente Portaria/Gestão/Admin registra presença.'),('RN-07','Notificação é criada para cada responsável vinculado após presença.'),('RN-08','Mensagens respeitam a matriz de perfis permitidos e não podem ser lidas por terceiros.'),('RN-09','Feedback DEMO_SEED é ilustrativo e não constitui evidência de pesquisa.'),('RN-10','Biometria, NFC/RFID, catraca e geolocalização contínua não pertencem ao MVP.'),('RN-11','Dados reais somente podem ser utilizados após autorização, base legal, governança e segurança adequadas.'),('RN-12','O dia escolar é calculado no fuso configurado; padrão da demonstração: America/Sao_Paulo.')]
add_table(doc, ['ID','Regra'], rules, [2,15], 9.5)

add_heading(doc, '7 CASOS DE USO TEXTUAIS', 1)
ucs = [
('UC-01','Autenticar-se','Responsável, Portaria, Gestão','Usuário ativo informa credenciais válidas.','Sistema valida credenciais, cria sessão temporária e registra login.','Credencial inválida: 401; excesso de tentativas: 429.'),
('UC-02','Registrar entrada','Portaria/Gestão','Sessão válida; aluno ativo; token válido.','Operador informa token e ENTRADA; sistema valida sequência, grava presença, notifica responsáveis e audita.','Token inválido: 404; entrada já pendente: 409.'),
('UC-03','Registrar saída','Portaria/Gestão','Sessão válida; existe entrada no mesmo dia.','Operador informa token e SAIDA; sistema valida sequência, grava, notifica e audita.','Sem entrada anterior: 409.'),
('UC-04','Consultar histórico/relatório','Responsável/Gestão/Portaria conforme escopo','Sessão válida.','Sistema calcula o conjunto de alunos permitido e devolve somente registros desse conjunto.','Sem sessão: 401.'),
('UC-05','Cadastrar aluno','Gestão','Sessão válida e perfil de gestão.','Gestão informa nome, matrícula e turma; sistema valida e gera token único.','Dados inválidos: 400; matrícula duplicada: 409.'),
('UC-06','Vincular responsável','Gestão','Responsável e aluno ativos.','Gestão seleciona ambos; sistema adiciona o vínculo sem duplicidade.','Entidade inexistente/inativa: 404.'),
('UC-07','Consultar/marcar notificação','Responsável','Sessão válida.','Usuário lista suas notificações e pode marcar a própria como lida.','Notificação de terceiro: 404.'),
('UC-08','Enviar mensagem','Responsável/Portaria/Gestão','Sessão válida e destinatário permitido.','Sistema valida a matriz de perfis, grava mensagem e a exibe apenas aos participantes.','Destinatário não permitido: 403.'),
('UC-09','Consultar auditoria','Gestão','Sessão válida e perfil autorizado.','Sistema exibe eventos recentes com usuário, ação, entidade e horário.','Perfil não autorizado: 403.'),
('UC-10','Registrar feedback','Usuário autenticado','Participante executou uma tarefa de validação.','Sistema registra perfil, cenário, sucesso, tempo, nota e comentário como APRESENTACAO.','Campos obrigatórios inválidos: 400.'),
('UC-11','Exportar validação','Gestão','Sessão válida e perfil de gestão.','Sistema gera CSV somente com linhas APRESENTACAO.','Perfil não autorizado: 403.'),
('UC-12','Restaurar demo','Gestão','Sessão válida e perfil de gestão.','Sistema repõe a seed, registra auditoria e invalida sessões.','Perfil não autorizado: 403.')]
add_table(doc, ['ID','Caso','Ator','Pré-condição','Fluxo principal','Alternativa/erro'], ucs, [1.2,3.2,3,4,6,4], 7.7)

sec = doc.add_section(WD_SECTION.NEW_PAGE); set_landscape(sec); header_footer(sec)
add_heading(doc, '8 DIAGRAMAS DE CASOS DE USO', 1)
add_body(doc, 'Para eliminar cruzamentos de linhas e relações artificiais, os casos de uso foram separados por ator. Em UML, a associação entre ator e caso de uso é representada por linha sólida simples, sem ponta de seta. Relações «include» e «extend» não foram usadas porque o modelo não precisa delas para expressar o comportamento do MVP.')
add_image(doc, '01_uc_responsavel.png', 'Figura 1 - Casos de uso do ator Responsável.', 8.9)
add_image(doc, '02_uc_portaria.png', 'Figura 2 - Casos de uso do ator Portaria.', 8.9)
add_image(doc, '03_uc_gestao.png', 'Figura 3 - Casos de uso do ator Gestão Escolar.', 8.9)

add_heading(doc, '9 DIAGRAMAS DE CLASSES', 1)
add_body(doc, 'O modelo de classes é conceitual e acompanha as estruturas efetivamente presentes no JSON e no código. Não foram criadas subclasses Responsável, Portaria e Gestão porque a implementação usa uma única entidade Usuario com o atributo perfil. Também não foram adicionados serviços, bancos ou integrações que não existam no MVP.')
add_image(doc, '04_classes_nucleo.png', 'Figura 4 - Classes do núcleo de presença.', 9.6)
add_image(doc, '05_classes_suporte.png', 'Figura 5 - Classes de comunicação, auditoria e validação.', 9.4)

add_heading(doc, '10 MODELO RELACIONAL (DER)', 1)
add_body(doc, 'O DER representa uma proposta lógica para substituir a persistência JSON quando houver evolução para banco relacional. Chaves, unicidades e relacionamentos derivam das estruturas atuais do MVP. O relacionamento entre responsáveis e alunos é N:N e, por isso, é resolvido pela tabela associativa guardian_students. As linhas representam relacionamentos e não setas de fluxo.')
add_image(doc, '06_der_nucleo.png', 'Figura 6 - DER lógico do núcleo de usuários, alunos, turmas, presença e notificações.', 9.8)
add_image(doc, '07_der_suporte.png', 'Figura 7 - DER lógico de mensagens, auditoria e feedback.', 9.4)

add_heading(doc, '11 DIAGRAMA DE SEQUÊNCIA', 1)
add_body(doc, 'O diagrama de sequência usa setas porque, nesse tipo de UML, a seta representa uma mensagem dirigida entre participantes. As linhas de retorno são tracejadas. O fluxo abaixo corresponde à rota POST /api/attendance e às funções de domínio utilizadas pelo backend.')
add_image(doc, '08_sequencia_presenca.png', 'Figura 8 - Sequência de registro de entrada/saída.', 9.5)

add_heading(doc, '12 ARQUITETURA IMPLEMENTADA', 1)
add_image(doc, '09_arquitetura_componentes.png', 'Figura 9 - Componentes existentes no MVP.', 9.3)
add_body(doc, 'O frontend é HTML/CSS/JavaScript; o backend usa Node.js e módulos nativos; regras reutilizáveis estão em lib/domain.js e lib/security.js; a demonstração persiste dados em JSON local. PostgreSQL/MySQL, Redis, OAuth2, Docker, push real e hardware são possibilidades futuras e não fazem parte da arquitetura implementada.')

sec = doc.add_section(WD_SECTION.NEW_PAGE); set_portrait(sec); header_footer(sec)
add_heading(doc, '13 RASTREABILIDADE', 1)
trace = [
('RF-01','POST /api/login; verifyPassword','api.test.js; security.test.js'),('RF-02','canRegisterAttendance/canManageSchool/canViewAudit; guards nas rotas','api.test.js; privacy-regression.test.js'),('RF-03','POST /api/students; uniqueStudentToken','api.test.js + inspeção'),('RF-04','POST /api/links','api.test.js + fluxo funcional'),('RF-05/RF-06','POST /api/attendance; validateAttendanceSequence','domain.test.js; api.test.js'),('RF-07/RF-13','notifications; PATCH /api/notifications/:id','api.test.js + fluxo funcional'),('RF-08/RF-09','allowedStudentIds; GET /api/dashboard; GET /api/reports','domain.test.js; api.test.js'),('RF-10','GET /api/reports.csv','api.test.js'),('RF-11','canMessageRole; visibleDirectory; GET/POST /api/messages','domain.test.js; api.test.js; privacy-regression.test.js'),('RF-12','audit(); GET /api/audit','privacy-regression.test.js + fluxo funcional'),('RF-14','GET/POST /api/feedback; GET /api/feedback.csv','api.test.js; privacy-regression.test.js'),('RF-15','POST /api/demo/reset','privacy-regression.test.js')]
add_table(doc, ['Requisito','Implementação relacionada','Evidência/teste'], trace, [3,8,6], 9)

add_heading(doc, '14 RESULTADO DA REVISÃO TÉCNICA', 1)
add_body(doc, 'A versão anterior misturava requisitos de negócio, itens de interface, metas futuras e afirmações sem critério de medição. Os diagramas também concentravam relações demais em uma única figura, o que favorecia linhas cruzadas e símbolos que pareciam setas de fluxo onde UML exigia apenas associação. A V2.0 corrige esse problema por separação de responsabilidade, multiplicidades explícitas, uso de linhas sem seta em associações e documentação da diferença entre UML, DER e sequência.')
add_table(doc, ['Problema identificado','Correção aplicada'], [
('Casos de uso com excesso de cruzamentos e associação visual ambígua.','Diagramas separados por ator, apenas associações sólidas sem seta.'),
('Classes de perfis que não existem no código.','Um único Usuario com atributo perfil, coerente com a implementação.'),
('DER com aparência de fluxograma.','Modelo lógico com PK/FK/UQ e cardinalidades; relacionamentos não usam setas.'),
('OWASP MASVS em aplicação web.','Referência corrigida para OWASP ASVS 5.0.0.'),
('RF de “painel de privacidade” tratado como requisito de negócio.','Reclassificado como recurso de interface/apoio, não como função essencial.'),
('RNF genéricos sem evidência.','RNFs reescritos com critério verificável e evidência indicada; desempenho permanece meta a medir.'),
('Gestão podendo visualizar mensagens/notificações de terceiros.','Backend revisado para escopo do destinatário/participante.'),
('Dados mutáveis versionados no db.json.','Base de execução passa a ser db.runtime.json, gerada a partir de db.seed.json e ignorada pelo Git.'),
('Feedback ilustrativo misturado na exportação.','CSV de validação passa a exportar apenas fonte APRESENTACAO.')], [7,10], 9.5)

add_heading(doc, 'REFERÊNCIAS', 1)
for ref in [
    'OBJECT MANAGEMENT GROUP. Unified Modeling Language (UML), Version 2.5.1. 2017.',
    'OWASP FOUNDATION. Application Security Verification Standard (ASVS) 5.0.0. 2025.',
    'WORLD WIDE WEB CONSORTIUM. Web Content Accessibility Guidelines (WCAG) 2.2. W3C Recommendation.',
    'BRASIL. Lei nº 13.709, de 14 de agosto de 2018. Lei Geral de Proteção de Dados Pessoais.',
    'ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS. NBR 6023: Informação e documentação - Referências - Elaboração. Edição vigente adotada pela instituição.',
    'ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS. NBR 10520: Informação e documentação - Citações em documentos - Apresentação. 2023.']:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8); p.paragraph_format.line_spacing = 1.0; p.add_run(ref)

doc.save(OUT)
print(f'Gerado: {OUT}')
