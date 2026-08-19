from pathlib import Path
from docx import Document
from docx.shared import Pt

path = Path('01_Plano_Gerenciamento_Unificado_Safe_Student.docx')
doc = Document(path)

research = {
    'Participantes da validação': 'Resultados históricos declarados na versão anterior; a evidência primária não foi localizada no repositório. Não apresentar quantitativo como validado sem anexar formulários ou CSV verificável.',
    'Tarefa 1 - localizar histórico': 'Pendente de comprovação primária ou de nova coleta controlada.',
    'Tarefa 2 - registrar presença': 'Pendente de comprovação primária ou de nova coleta controlada.',
    'Tarefa 3 - compreender notificação': 'Pendente de comprovação primária ou de nova coleta controlada.',
    'Satisfação geral': 'Pendente de comprovação primária. A nova coleta deve registrar nota individual no módulo Feedback.',
    'Principal melhoria sugerida': 'A nova coleta deve registrar cenário, conclusão da tarefa, tempo, nota e comentário; somente dados APRESENTACAO podem ser tratados como evidência coletada.',
}

for table in doc.tables:
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        key = row.cells[0].text.strip()
        if key in research:
            row.cells[1].text = research[key]
            for p in row.cells[1].paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'; r.font.size = Pt(9)

replacements = [
    ('Versão | V1.1', 'Versão | V2.0'),
    ('Versão | V1', 'Versão | V2.0'),
    ('15 RF/RNF', 'requisitos funcionais e não funcionais rastreáveis'),
    ('Validar 5 funcionalidades essenciais e 15 RF/RNF com testes automatizados e avaliação acadêmica controlada de usabilidade.', 'Validar as 5 funcionalidades essenciais por requisitos rastreáveis, testes automatizados e avaliação acadêmica controlada de usabilidade.'),
]
for p in doc.paragraphs:
    full = ''.join(r.text for r in p.runs)
    new = full
    for old, repl in replacements:
        new = new.replace(old, repl)
    if new != full:
        p.text = new
        for r in p.runs:
            r.font.name = 'Times New Roman'; r.font.size = Pt(12)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                full = ''.join(r.text for r in p.runs); new = full
                for old, repl in replacements:
                    new = new.replace(old, repl)
                if new != full:
                    p.text = new
                    for r in p.runs:
                        r.font.name = 'Times New Roman'; r.font.size = Pt(9)

doc.save(path)
print('Plano V2 sincronizado sem fabricar evidência de pesquisa.')
