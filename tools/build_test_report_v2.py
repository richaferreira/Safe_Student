from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('04_Plano_Relatorio_Testes_Safe_Student.docx')

def shade(cell, fill='D9E2F3'):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def table(doc, headers, rows, font=9):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers): t.rows[0].cells[i].text = h; shade(t.rows[0].cells[i])
    for row in rows:
        c = t.add_row().cells
        for i,v in enumerate(row): c[i].text = str(v)
    for ri,row in enumerate(t.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name='Times New Roman'; r.font.size=Pt(font); r.bold=(ri==0)
    return t

def body(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.first_line_indent=Cm(1.25); p.paragraph_format.line_spacing=1.5; p.add_run(text)

def heading(doc,text,level=1): doc.add_paragraph(text,style=f'Heading {level}')

doc=Document(); sec=doc.sections[0]; sec.top_margin=Cm(3); sec.left_margin=Cm(3); sec.right_margin=Cm(2); sec.bottom_margin=Cm(2)
doc.styles['Normal'].font.name='Times New Roman'; doc.styles['Normal'].font.size=Pt(12)
for level in (1,2):
    st=doc.styles[f'Heading {level}']; st.font.name='Times New Roman'; st.font.bold=True; st.font.size=Pt(14 if level==1 else 12)
for _ in range(3): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('SAFE STUDENT'); r.bold=True; r.font.size=Pt(18)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('PLANO E RELATÓRIO DE TESTES — V2.0'); r.bold=True; r.font.size=Pt(15)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Evidências automatizadas do MVP acadêmico após auditoria sênior')
doc.add_page_break()

heading(doc,'1 ESTRATÉGIA')
body(doc,'A estratégia V2 verifica domínio, segurança, API HTTP e regressões de privacidade. O objetivo não é aumentar artificialmente a quantidade de casos, e sim demonstrar regras importantes: autenticação, autorização por perfil, sequência de presença, isolamento de dados, exportações autenticadas, validação acadêmica e restauração controlada da demonstração.')

heading(doc,'2 AMBIENTE')
table(doc,['Item','Configuração'],[
('Runtime local de revisão','Node.js 22'),('Compatibilidade declarada','Node.js 18+'),('Comando','npm test'),('Framework','node:test'),('Persistência dos testes','arquivo JSON temporário isolado'),('CI','GitHub Actions configurado para Node 18, 20 e 22; confirmar estado antes da entrega final')],9.5)

heading(doc,'3 RESULTADO LOCAL DA REVISÃO V2')
body(doc,'A suíte foi executada localmente durante a revisão sênior: 24 testes aprovados, 0 falhas. Esse resultado é evidência da execução local da revisão e não deve ser confundido com o status do GitHub Actions; o CI deve ser conferido separadamente.')

cases=[
('UT-01','Normalização do token','PASS'),('UT-02','RBAC de presença','PASS'),('UT-03','RBAC de mensagens','PASS'),('UT-04','Responsável vê somente alunos vinculados','PASS'),('UT-05','Bloqueio de saída sem entrada','PASS'),('UT-06','Bloqueio de segunda entrada','PASS'),('UT-07','Saída válida após entrada','PASS'),('UT-08','Taxa usa chave do dia escolar fornecida','PASS'),('UT-09','Hash/verificação de senha','PASS'),('UT-10','Hash inválido falha com segurança','PASS'),('UT-11','Token aleatório possui tamanho/entropia esperados','PASS'),
('API-01','Health check sem autenticação','PASS'),('API-02','Login inválido retorna 401','PASS'),('API-03','Responsável não registra presença','PASS'),('API-04','Portaria registra entrada e segunda entrada é bloqueada','PASS'),('API-05','Destinatário indevido é bloqueado pela API','PASS'),('API-06','CSV de relatório exige autenticação','PASS'),('API-07','Feedback separa DEMO_SEED de APRESENTACAO','PASS'),('API-08','Somente gestão exporta validação','PASS'),
('PRIV-01','Diretório expõe somente dados mínimos','PASS'),('PRIV-02','Gestão não lê mensagem privada de terceiros','PASS'),('PRIV-03','Gestão não recebe notificação destinada ao responsável','PASS'),('PRIV-04','CSV de validação exclui DEMO_SEED','PASS'),('PRIV-05','Restauração da demo gera evento de auditoria','PASS')]
table(doc,['ID','Caso automatizado','Resultado'],cases,9)

heading(doc,'4 CASOS FUNCIONAIS PARA A BANCA')
table(doc,['ID','Roteiro','Resultado esperado'],[
('BF-01','Login como Portaria e registrar ENTRADA com SS-ALU001','Registro criado; responsáveis vinculados recebem notificação interna.'),('BF-02','Tentar nova ENTRADA para o mesmo aluno no mesmo dia','Operação recusada por sequência inválida.'),('BF-03','Login como Responsável','Somente alunos vinculados e suas informações permitidas são exibidos.'),('BF-04','Exportar relatório do Responsável','CSV autenticado contém apenas escopo permitido.'),('BF-05','Tentar mensagem para destinatário não permitido por chamada direta','Backend retorna 403.'),('BF-06','Login como Gestão e consultar auditoria','Eventos críticos aparecem com usuário, ação e data/hora.'),('BF-07','Registrar feedback real e exportar validação','CSV contém APRESENTACAO e não contém DEMO_SEED.'),('BF-08','Restaurar a demonstração','Seed é restaurada, ação auditada e sessão atual invalidada.')],8.8)

heading(doc,'5 LIMITAÇÕES DOS TESTES')
body(doc,'A suíte não substitui teste de carga, pentest, homologação de acessibilidade, avaliação jurídica/LGPD, teste de integração com hardware ou serviços externos. Esses itens pertencem a uma eventual evolução para produção. O RNF de desempenho de até 2 segundos permanece uma meta a medir com procedimento reproduzível antes de ser apresentado como atendido.')

heading(doc,'6 CRITÉRIOS DE ACEITE')
table(doc,['Critério','Situação V2'],[
('Sintaxe do backend/frontend','Deve ser validada no CI e/ou localmente antes da apresentação.'),('Suíte automatizada','24/24 na execução local da revisão.'),('Autorização server-side','Coberta por testes de API e privacidade.'),('Dados de pesquisa','Parcial: somente coleta real pode fechar a evidência primária.'),('MVP de produção','Não aplicável: projeto é demonstração acadêmica.')],9.3)

doc.save(OUT)
print(f'Gerado: {OUT}')
