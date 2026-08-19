from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def setup_doc(title, subtitle):
    doc=Document(); sec=doc.sections[0]; sec.top_margin=Cm(3); sec.left_margin=Cm(3); sec.right_margin=Cm(2); sec.bottom_margin=Cm(2)
    doc.styles['Normal'].font.name='Times New Roman'; doc.styles['Normal'].font.size=Pt(12)
    for level in (1,2):
        st=doc.styles[f'Heading {level}']; st.font.name='Times New Roman'; st.font.bold=True; st.font.size=Pt(14 if level==1 else 12)
    for _ in range(3): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('SAFE STUDENT'); r.bold=True; r.font.size=Pt(18)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(title); r.bold=True; r.font.size=Pt(15)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(subtitle)
    doc.add_page_break(); return doc

def shade(cell,fill='D9E2F3'):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def table(doc,headers,rows,font=9):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers): t.rows[0].cells[i].text=h; shade(t.rows[0].cells[i])
    for row in rows:
        c=t.add_row().cells
        for i,v in enumerate(row): c[i].text=str(v)
    for ri,row in enumerate(t.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after=Pt(0)
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(font); r.bold=(ri==0)
    return t

def body(doc,text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.first_line_indent=Cm(1.25); p.paragraph_format.line_spacing=1.5; p.add_run(text)

def heading(doc,text,level=1): doc.add_paragraph(text,style=f'Heading {level}')

# 06 Auditoria
adoc=setup_doc('RELATÓRIO DE AUDITORIA TÉCNICA E DE CONFORMIDADE — V2.0','Revisão sênior de requisitos, UML/DER, código, privacidade, testes e evidências')
heading(adoc,'1 OBJETIVO')
body(adoc,'A auditoria V2.0 confronta cada afirmação técnica com o comportamento observável do MVP. A revisão evita tratar metas futuras como funcionalidades implementadas e evita tratar resultados históricos de pesquisa como evidência auditada quando os registros primários não estão disponíveis.')
heading(adoc,'2 ACHADOS E CORREÇÕES')
table(adoc,['Severidade','Achado','Correção V2'],[
('Crítico','Casos de uso com cruzamentos e linhas que pareciam setas de fluxo.','Diagramas separados por ator; associação sólida simples, sem seta.'),('Crítico','Classes misturavam conceito, perfil e estrutura relacional.','Modelo conceitual separado do DER; Usuario único com PerfilUsuario e multiplicidades.'),('Crítico','DER parecia fluxograma.','DER lógico com PK/FK/UQ e cardinalidades, sem setas de processo.'),('Alto','Referência OWASP MASVS em aplicação web.','OWASP ASVS 5.0.0 usado como referência de segurança web.'),('Alto','Gestão podia visualizar mensagens/notificações privadas de terceiros.','Escopo restrito aos participantes/destinatários; testes de regressão.'),('Alto','Diretório devolvia dados pessoais além do necessário.','Projeção mínima: id, nome, perfil e status.'),('Alto','CSV de validação incluía DEMO_SEED.','Exportação somente de APRESENTACAO.'),('Médio','Taxa de presença usava corte UTC diferente do dia escolar.','Cálculo unificado pela função dateKey do fuso escolar.'),('Médio','Token de novo aluno não verificava colisão explicitamente.','Geração com verificação de unicidade.'),('Médio','Base mutável db.json era versionada.','db.runtime.json é criado da seed e ignorado pelo Git.'),('Acadêmico','Números históricos de pesquisa apareciam como resultado sem evidência primária localizada.','Status da pesquisa permanece Parcial; nova coleta real é necessária.')],8.6)
heading(adoc,'3 REQUISITOS E MODELAGEM')
body(adoc,'Os requisitos funcionais foram reescritos como obrigações observáveis com critério de aceite. Os requisitos não funcionais indicam como devem ser comprovados. A modelagem usa associação sem seta em casos de uso e classes; o DER apresenta relacionamentos com chaves e cardinalidades; o diagrama de sequência usa setas porque elas representam mensagens; o diagrama de componentes mostra somente a arquitetura efetivamente implementada.')
heading(adoc,'4 AUDITORIA DO CÓDIGO')
table(adoc,['Área','Conclusão'],[
('server.js','Autenticação, autorização, minimização, presença, mensagens, notificações, relatórios, feedback e reset revisados.'),('lib/domain.js','RBAC, escopo de alunos, sequência de presença e taxa por dia escolar revisados.'),('lib/security.js','scrypt + timingSafeEqual adequados ao MVP; não são apresentados como arquitetura de alta carga.'),('public/','Interface usa API autenticada e escape de conteúdo; controles visuais não substituem autorização server-side.'),('data/','Seed sintética versionada; runtime mutável não deve ser commitado.'),('tests/','Domínio, segurança, API e regressões de privacidade.')],9)
heading(adoc,'5 TESTES')
body(adoc,'Na revisão sênior, a suíte V2 foi executada localmente em Node.js 22 com 24 testes aprovados e 0 falhas. O workflow de CI permanece configurado para Node.js 18, 20 e 22 e deve ser conferido separadamente antes da entrega final; o resultado local não deve ser descrito como se fosse um status do GitHub Actions.')
heading(adoc,'6 EVIDÊNCIA EXTENSIONISTA')
body(adoc,'A pesquisa de campo continua Parcial até existir evidência primária real: formulários/termos preenchidos, respostas anonimizadas ou CSV de coleta que permita reproduzir os resultados. Não se deve inventar participantes, respostas, percentuais ou médias. Registros DEMO_SEED são ilustrativos e não constituem pesquisa.')
heading(adoc,'7 SITUAÇÃO V2')
table(adoc,['Dimensão','Status'],[
('Código/MVP','Revisado; suíte local 24/24.'),('Requisitos','Reescritos e rastreáveis.'),('UML/DER','Refazidos com notação formal e responsabilidades separadas.'),('Documentação','Sincronizada com a revisão V2 nos artefatos finais.'),('Pesquisa de campo','Parcial até evidência primária real.'),('Produção','Fora do escopo: MVP acadêmico local, não homologado.')],9.5)
heading(adoc,'REFERÊNCIAS')
for ref in ['OBJECT MANAGEMENT GROUP. Unified Modeling Language (UML), Version 2.5.1.','OWASP FOUNDATION. Application Security Verification Standard (ASVS) 5.0.0.','WORLD WIDE WEB CONSORTIUM. Web Content Accessibility Guidelines (WCAG) 2.2.','BRASIL. Lei nº 13.709/2018 — Lei Geral de Proteção de Dados Pessoais.']:
    adoc.add_paragraph(ref)
adoc.save('06_Relatorio_Auditoria_Conformidade_Safe_Student.docx')

# 10 Manual
mdoc=setup_doc('MANUAL DO USUÁRIO, DEMONSTRAÇÃO E SUPORTE — V2.0','Orientações alinhadas ao comportamento real do MVP revisado')
heading(mdoc,'1 VISÃO GERAL')
body(mdoc,'Este manual orienta a execução e a demonstração do Safe Student V2. O sistema é um MVP local com dados sintéticos e perfis de Responsável, Portaria e Gestão Escolar. Não deve receber dados reais sem autorização institucional, governança e segurança adequadas.')
heading(mdoc,'2 INSTALAÇÃO E EXECUÇÃO')
table(mdoc,['Passo','Ação'],[
('1','Instalar Node.js 18 ou superior.'),('2','Abrir terminal na pasta Safe_Student_MVP.'),('3','Executar npm start.'),('4','Acessar http://localhost:3000.'),('5','Opcional: executar npm test.'),('6','Para restaurar o estado local pelo terminal, executar npm run reset.')],9.5)
body(mdoc,'A aplicação usa data/db.seed.json como estado inicial sintético. Na primeira execução é criado data/db.runtime.json, que contém o estado mutável e não deve ser versionado.')
heading(mdoc,'3 PERFIS DE DEMONSTRAÇÃO')
table(mdoc,['Perfil','E-mail','Senha','Uso principal'],[
('Responsável','responsavel@demo.com','demo123','Alunos vinculados, notificações, histórico, relatório, mensagens e feedback.'),('Portaria','portaria@demo.com','demo123','Registro de entrada/saída e comunicação permitida.'),('Gestão','gestor@demo.com','demo123','Cadastros, vínculos, relatórios, auditoria, validação e reset.')],8.8)
heading(mdoc,'4 FLUXO RECOMENDADO PARA A BANCA')
table(mdoc,['Etapa','Procedimento'],[
('1','Entrar como Portaria e registrar ENTRADA com SS-ALU001.'),('2','Mostrar que uma segunda ENTRADA é bloqueada.'),('3','Entrar como Responsável e mostrar a notificação e o histórico do aluno vinculado.'),('4','Exportar o relatório autenticado.'),('5','Demonstrar uma mensagem permitida e explicar que a autorização também existe no backend.'),('6','Entrar como Gestão e mostrar auditoria.'),('7','Registrar feedback real, exportar validação e explicar que DEMO_SEED não é evidência.'),('8','Restaurar a demonstração ao final, se necessário.')],9)
heading(mdoc,'5 SEGURANÇA E PRIVACIDADE NA DEMONSTRAÇÃO')
for item in ['Use somente os dados sintéticos fornecidos.', 'Não demonstre acesso indevido a mensagens/notificações de terceiros como se fosse recurso de gestão.', 'Não apresente o MVP como certificado LGPD/WCAG/OWASP.', 'Explique que biometria, NFC/RFID, catracas, geolocalização e push real estão fora do escopo.', 'A pesquisa de campo só pode ser considerada concluída com evidência primária real.']:
    p=mdoc.add_paragraph(style='List Bullet'); p.add_run(item)
heading(mdoc,'6 TESTES E SUPORTE')
body(mdoc,'Execute npm test antes da apresentação. Na revisão V2 foram aprovados localmente 24/24 testes em Node.js 22. O CI está configurado para Node.js 18, 20 e 22 e deve ser conferido no GitHub. Problemas comuns incluem porta 3000 ocupada, Node.js ausente, base local alterada ou credencial digitada incorretamente.')
heading(mdoc,'7 LIMITAÇÕES')
body(mdoc,'O MVP usa JSON local e sessões em memória. Não possui HTTPS próprio, banco transacional, backup corporativo, SSO, observabilidade, SMS/push real ou integração com hardware. Uma implantação real exigiria arquitetura e governança adicionais.')
mdoc.save('10_Manual_Usuario_Implantacao_Treinamento_Suporte.docx')
print('Auditoria e manual V2 gerados.')
